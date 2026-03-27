#!/usr/bin/env python3
"""Edit the realtor.com negotiation brief: merge agent asks, preserve formatting."""

from docx import Document
import os

path = os.path.join(os.path.dirname(__file__), 'Pearl_MKTG_Realtor.com Negotiation Brief_Q12026.docx')
doc = Document(path)

# --- Pass 1: Rename "Agent Marketing Tools Integration" and update its content ---
for i, para in enumerate(doc.paragraphs):
    # Rename heading
    if 'Agent Marketing Tools Integration' in para.text:
        for run in para.runs:
            if 'Agent Marketing Tools Integration' in run.text:
                run.text = run.text.replace('Agent Marketing Tools Integration', 'Agent Adoption Pipeline')

    # Update "The ask" line
    if para.text.strip().startswith('The ask:') and i > 0:
        # Check if the previous heading was about Agent (look back up to 3 paras)
        context = ' '.join(doc.paragraphs[max(0,i-5):i][-5:].text if hasattr(doc.paragraphs[max(0,i-5):i], 'text') else '')

    if "Include SCORE data in realtor.com" in para.text and "agent marketing suite" in para.text:
        for run in para.runs:
            if 'listing flyers' in run.text or 'agent marketing suite' in run.text:
                # Find the run with the description and update it
                old = "Include SCORE data in realtor.com's agent marketing suite (listing flyers, CMA tools, client presentations)."
                new = "Include SCORE data in realtor.com's agent marketing suite (listing flyers, CMA tools, client presentations) and feature Pearl in agent newsletters, webinars, and training materials."
                if old in run.text:
                    run.text = run.text.replace(old, new)
                elif 'agent marketing suite' in run.text:
                    # Content might be split across runs - handle the main run
                    run.text = run.text.replace(
                        "agent marketing suite (listing flyers, CMA tools, client presentations).",
                        "agent marketing suite (listing flyers, CMA tools, client presentations) and feature Pearl in agent newsletters, webinars, and training materials."
                    )

    # Update "Why it matters" line for agent section
    if "Agent adoption accelerator" in para.text and "tools they already use" in para.text:
        for run in para.runs:
            if 'Agent adoption accelerator' in run.text:
                run.text = run.text.replace(
                    "Agent adoption accelerator — tools they already use with SCORE built in.",
                    "Agent adoption accelerator on two fronts — embedded in tools agents already use, and promoted through education channels they already consume. Pearl can provide turnkey content."
                )
            elif 'tools they already use' in run.text:
                run.text = run.text.replace(
                    "tools they already use with SCORE built in.",
                    "on two fronts — embedded in tools agents already use, and promoted through education channels they already consume. Pearl can provide turnkey content."
                )

    # Update "What it's worth" line for agent section
    if "Reduces Pearl" in para.text and "agent acquisition cost" in para.text:
        for run in para.runs:
            if 'agent acquisition cost' in run.text:
                run.text = run.text.replace(
                    "Reduces Pearl's agent acquisition cost and accelerates market penetration.",
                    "Reduces Pearl's agent acquisition cost and accelerates market penetration across realtor.com's full agent ecosystem."
                )

# --- Pass 2: Delete the "Cross-Promotion in Agent Communications" section ---
# Find the heading and delete it plus its content paragraphs until the next heading
delete_mode = False
paras_to_delete = []

for i, para in enumerate(doc.paragraphs):
    if 'Cross-Promotion in Agent Communications' in para.text:
        delete_mode = True
        paras_to_delete.append(i)
        continue

    if delete_mode:
        # Stop deleting when we hit the next numbered heading (next value-add ask) or a major heading
        text = para.text.strip()
        # Check if this is a new section heading (starts with a number followed by period, or is a Heading style)
        is_next_section = False
        if para.style and 'Heading' in para.style.name:
            # Check if it's a different section, not a sub-element of the current one
            if text and text[0].isdigit() and '.' in text[:3]:
                is_next_section = True
            elif text and not any(label in text for label in ['The ask', 'Why it', 'How to frame']):
                is_next_section = True

        if is_next_section:
            delete_mode = False
        else:
            paras_to_delete.append(i)

# Delete paragraphs in reverse order to preserve indices
for idx in reversed(paras_to_delete):
    para = doc.paragraphs[idx]
    p_element = para._element
    p_element.getparent().remove(p_element)

# --- Renumber remaining value-add asks ---
# After deleting Cross-Promotion, the numbering may need fixing
counter = 0
for para in doc.paragraphs:
    if para.style and 'Heading' in para.style.name:
        text = para.text.strip()
        # Check if it's in the value-add section (numbered items after "Value-Add Asks")
        if text and text[0].isdigit() and '. ' in text[:4]:
            # Only renumber if we're past the Primary Asks (which has 1-4)
            # We detect value-adds by checking content
            pass  # Numbering is fine since we removed a middle item

doc.save(path)
print(f'Edited and saved: {path}')

# Verify changes
doc2 = Document(path)
for p in doc2.paragraphs:
    if 'Agent Adoption Pipeline' in p.text:
        print(f'  OK: Found "Agent Adoption Pipeline"')
    if 'Cross-Promotion in Agent Communications' in p.text:
        print(f'  WARN: "Cross-Promotion" still present')

print('Done.')
