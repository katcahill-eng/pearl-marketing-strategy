#!/usr/bin/env python3
"""Generate a formatted .docx negotiation brief for realtor.com."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x4B)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)

# -- Title --
title = doc.add_heading('Negotiation Brief: Realtor.com Partnership', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = meta.add_run('Pearl Marketing  |  March 2026  |  CONFIDENTIAL')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.italic = True

doc.add_paragraph()  # spacer

# -- Partner Profile --
doc.add_heading('Partner Profile', level=2)

doc.add_paragraph(
    'Realtor.com is the second-largest US property search platform with ~100 million monthly unique visitors, '
    'trailing only Zillow. As the official website of NAR, they have unmatched credibility with real estate '
    'professionals and strong relationships with MLSs nationwide. Their audience includes both serious home '
    'shoppers (higher intent than Zillow browsers) and the majority of practicing real estate agents.'
)

doc.add_heading('Key Assets Pearl Wants', level=3)
assets = [
    'Consumer traffic during active home shopping (high-intent audience)',
    'Agent adoption pipeline through existing realtor relationships',
    'Content distribution through their editorial platform and agent marketing tools',
    'Data integration infrastructure already connected to MLSs nationwide',
    'NAR credibility halo that validates Pearl as a legitimate industry tool',
]
for a in assets:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('Their Strategic Goals', level=3)
goals = [
    'Differentiate from Zillow with unique data and tools',
    'Increase agent engagement and loyalty (defend against Zillow\'s agent recruitment)',
    'Build deeper consumer engagement beyond basic search',
    'Position as the most comprehensive property information platform',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

# -- Pearl's Leverage --
doc.add_heading("Pearl's Leverage", level=2)

doc.add_paragraph(
    'Pearl holds the only comprehensive national home performance dataset. No competitor has national-scale '
    'data across Pearl\'s five pillars (Safety, Comfort, Operations, Resilience, Energy). This positions Pearl '
    'as realtor.com\'s best opportunity to offer differentiated property intelligence that Zillow cannot match.'
)

leverage = [
    ('92+ million home performance profiles', 'largest dataset of its kind'),
    ('First-mover advantage', 'in home performance scoring with established methodology'),
    ('NAR partnership alignment', "Pearl's DOE and RESO relationships complement realtor.com's NAR affiliation"),
    ('Agent champion network', 'Pearl has agents already using SCORE who can evangelize to realtor.com users'),
    ('Future data moat', "Pearl's contractor certification network generates data streams competitors cannot access"),
    ('Content authority', "Pearl's thought leadership in home performance creates ongoing content opportunities"),
]
for title, desc in leverage:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title} — ')
    run.bold = True
    p.add_run(desc)

p = doc.add_paragraph()
run = p.add_run('The core leverage: ')
run.bold = True
p.add_run('Realtor.com needs unique data to compete with Zillow. Pearl has the most comprehensive and differentiated property data available.')

# -- Primary Asks --
doc.add_heading('Primary Asks', level=2)

primary_asks = [
    ('SCORE Integration in Property Details',
     'Display Pearl SCORE prominently on property detail pages for all listings.',
     'Primary consumer exposure and brand building at scale.',
     'Equivalent to $50M+ in consumer advertising spend based on realtor.com\'s traffic.'),
    ('Agent Adoption Pipeline',
     'Include SCORE data in realtor.com\'s agent marketing suite (listing flyers, CMA tools, client presentations) and feature Pearl in agent newsletters, webinars, and training materials.',
     'Agent adoption accelerator on two fronts — embedded in tools agents already use, and promoted through education channels they already consume. Pearl can provide turnkey content.',
     'Reduces Pearl\'s agent acquisition cost and accelerates market penetration across realtor.com\'s full agent ecosystem.'),
    ('Content Partnership and Thought Leadership Platform',
     'Regular home performance content placement in realtor.com\'s editorial sections and agent newsletters.',
     'Builds category authority and drives organic search rankings.',
     '$2M+ annually in equivalent PR and content marketing value.'),
    ('Data Licensing Revenue Share',
     'Revenue sharing model where Pearl receives percentage of any premium product sales that include SCORE data.',
     'Direct monetization of data integration, not just exposure.',
     'Potential $5–10M annually based on realtor.com\'s premium product revenue.'),
]

for i, (title, ask, why, worth) in enumerate(primary_asks, 1):
    doc.add_heading(f'{i}. {title}', level=3)
    for label, text in [('The ask', ask), ('Why it matters', why), ("What it's worth", worth)]:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        p.add_run(text)

# -- Value-Add Asks --
doc.add_heading('Value-Add Asks', level=2)

p = doc.add_paragraph()
run = p.add_run('If realtor.com pushes back on price, these are high-value items for Pearl that cost realtor.com very little.')
run.italic = True

value_adds = [
    ('Publish Pearl Research on Realtor.com Blog',
     'Publish Pearl-authored research and data analysis on realtor.com\'s blog and editorial sections.',
     'Authoritative SEO backlinks from a DA 92 domain; positions Pearl as a thought leader to realtor.com\'s audience.',
     'They need editorial content anyway; Pearl provides it ready-to-publish.',
     '"We have original research on home operating costs and performance trends that your audience would find valuable — and it fills your editorial calendar."'),
    ('Beta User Community Access',
     'Access to realtor.com\'s agent advisory panels and beta user groups for Pearl product testing.',
     'Direct agent feedback loop for product development and messaging.',
     'They already run these programs; adding Pearl topics costs nothing.',
     '"We\'d love your agents\' input to make sure SCORE integrates seamlessly with their workflow."'),
    ('Consumer Email List Segmentation',
     'Ability to send targeted home performance content to realtor.com users who viewed homes in specific SCORE ranges.',
     'Direct consumer marketing channel with high-intent audiences.',
     'Uses existing email infrastructure; Pearl provides the content.',
     '"This helps both of us deliver more relevant content to users based on the homes they\'re actually viewing."'),
    ('PR and Media Amplification',
     'Joint press releases for SCORE milestones, realtor.com quotes Pearl in industry media, inclusion in speaking opportunities.',
     'Credibility boost and earned media coverage.',
     'Uses existing PR relationships and speaking slots they already have.',
     '"Partnership announcements get coverage for both brands — let\'s maximize the PR value."'),
    ('Search and SEO Link Equity',
     'Strategic internal linking from realtor.com pages to Pearl content and resources.',
     'SEO authority transfer and organic search improvement.',
     'Minor editorial adjustment to existing content.',
     '"Cross-linking improves user experience — when someone views a home, they can easily learn more about performance."'),
    ('Mobile App Integration Testing',
     'A/B test SCORE integration in realtor.com\'s mobile app with select user segments.',
     'Mobile exposure and conversion data.',
     'Limited test scope, uses existing app development resources.',
     '"Let\'s test this with a small user group first to prove the value before full rollout."'),
    ('Agent Success Story Co-Creation',
     'Joint case studies featuring realtor.com agents who use Pearl SCORE successfully.',
     'Agent testimonials and use case documentation.',
     'They already create agent spotlight content.',
     '"Success stories help both of us — agents love recognition and it shows other agents what\'s possible."'),
]

for i, (title, ask, value, cost, frame) in enumerate(value_adds, 1):
    doc.add_heading(f'{i}. {title}', level=3)
    for label, text in [('The ask', ask), ("Why it's valuable to Pearl", value), ("Why it's low-cost for the partner", cost), ('How to frame it', frame)]:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        if label == 'How to frame it':
            run2 = p.add_run(text)
            run2.italic = True
        else:
            p.add_run(text)

# -- Walk-Away Criteria --
doc.add_heading('Walk-Away Criteria', level=2)

doc.add_heading('Hard Red Lines', level=3)
red_lines = [
    ('No exclusivity without premium pricing', 'if realtor.com demands exclusive SCORE data, Pearl must receive 3x standard licensing fees minimum'),
    ('No co-branding restrictions', 'Pearl must maintain independent brand identity in all integrations'),
    ('No data ownership transfer', 'realtor.com cannot claim ownership of Pearl data or methodology'),
    ('No agent fee requirements', 'realtor.com cannot require agents to pay additional fees to access SCORE data'),
]
for title, desc in red_lines:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title} — ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('Soft Red Lines', level=3)
doc.add_paragraph('Revenue sharing below 15% — acceptable but seek other value-adds to compensate', style='List Bullet')
doc.add_paragraph('Limited integration scope — partial integration is better than none, but push for expansion timeline', style='List Bullet')
doc.add_paragraph('Content restrictions — Pearl needs ability to create educational content, but can accept editorial guidelines', style='List Bullet')

doc.add_heading('Deal Quality Assessment', level=3)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.LEFT
cells = table.rows[0].cells
cells[0].text = 'Quality'
cells[1].text = 'What It Looks Like'
for cell in cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
data = [
    ('Great deal', 'Full integration + revenue share + marketing amplification'),
    ('Good deal', 'Full integration + significant marketing support'),
    ('Acceptable', 'Limited integration + some marketing support + clear expansion path'),
    ('Bad deal', 'Data access only with no marketing support and restrictive terms'),
]
for i, (q, d) in enumerate(data, 1):
    table.rows[i].cells[0].text = q
    table.rows[i].cells[1].text = d

# -- Negotiation Sequence --
doc.add_heading('Negotiation Sequence', level=2)

for label, items in [
    ('Open with', ['Pearl\'s unique market position and data differentiation opportunity for realtor.com. Lead with their competitive need against Zillow rather than Pearl\'s need for distribution.']),
    ('Trade early', [
        'Exclusive preview period (60–90 days before other platforms get SCORE)',
        'Co-marketing budget allocation for joint promotions',
        'Technical integration timeline flexibility',
    ]),
    ('Hold back', [
        'Agent training and onboarding support (Pearl can provide but frame as additional value)',
        'Custom content creation (offer as sweetener if negotiations stall)',
        'Enhanced API access with real-time updates (premium tier offering)',
    ]),
    ('Escalation path', [
        'If price negotiations stall → introduce value-add asks to increase deal value without changing base price',
        'If integration scope gets limited → push for pilot program with expansion milestones',
        'If decision timeline extends → create urgency around competitive advantage window',
    ]),
]:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    if len(items) == 1:
        p.add_run(items[0])
    else:
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Framework: ')
run.bold = True
p.add_run('Use value stacking — layer multiple elements so the total deal value exceeds the sum of parts. Never negotiate price alone; always negotiate value packages.')

# -- Talking Points --
doc.add_heading('Talking Points', level=2)

talking_points = [
    ('Primary Value Proposition', 'Realtor.com can offer the most comprehensive property intelligence in the market — Zillow has Zestimate for price, but you\'ll have SCORE for performance, which is what buyers actually live with every day.'),
    ('Data Differentiation', 'Pearl has the only national database of home performance — this isn\'t something realtor.com can build internally, and it\'s not available anywhere else at this scale.'),
    ('Agent Win-Win', 'Your agents get a competitive advantage tool, Pearl gets adoption, and realtor.com gets agent loyalty — everyone wins when agents can offer buyers something unique.'),
    ('Consumer Value', 'Home buyers spend $21,400 per year beyond their mortgage on operating costs — SCORE helps them make smarter decisions upfront instead of discovering problems after closing.'),
    ('Market Timing', 'Home performance is becoming mainstream thanks to IRA incentives and climate awareness — realtor.com can lead this trend rather than follow it.'),
    ('Partnership Precedent', 'We\'re working with NAR and DOE on industry standards — this positions realtor.com alongside the official industry bodies rather than outside them.'),
]

for title, point in talking_points:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run2 = p.add_run(f'"{point}"')
    run2.italic = True

# -- Footer --
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Prepared by Pearl Marketing OS  |  March 2026  |  CONFIDENTIAL')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.italic = True

# Save
output_path = os.path.join(os.path.dirname(__file__), 'Pearl_MKTG_Realtor.com_Negotiation_Brief_Q12026.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
