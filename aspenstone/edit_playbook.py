"""
Edit AspenStone SEO+GEO Playbook v4 in place.

Reads /tmp/aspen.orig.docx, applies the restructure described in the handoff
(ads-driven Phase 2, new budget, new city-page schedule, etc.), and writes
the result to aspenstone/AspenStone_SEO_GEO_Playbook_v5.docx in the repo.
"""
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

SRC = "/tmp/aspen.orig.docx"
DST = "/home/user/pearl-marketing-strategy/aspenstone/AspenStone_SEO_GEO_Playbook_v5.docx"


# ---------- helpers ----------

def set_text(paragraph, new_text):
    """Replace paragraph text, keeping the formatting of the first run."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def set_cell_text(cell, new_text):
    """Replace a table cell's text, keeping its first paragraph's first-run formatting."""
    # Remove all but the first paragraph
    paragraphs = cell.paragraphs
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    set_text(paragraphs[0], new_text)


def set_cell_multi(cell, lines):
    """Replace a cell with multiple lines, each as its own paragraph cloned from the first."""
    paragraphs = cell.paragraphs
    template = paragraphs[0]
    # Remove extra paragraphs
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    set_text(template, lines[0])
    # Append additional paragraphs cloned from the template
    parent = template._element.getparent()
    last = template._element
    for line in lines[1:]:
        clone = deepcopy(template._element)
        last.addnext(clone)
        last = clone
        # Set the text on the cloned paragraph
        # Easier: walk runs of the new clone
        from docx.text.paragraph import Paragraph
        p_new = Paragraph(clone, template._parent)
        set_text(p_new, line)


def clone_after(template_para, new_text):
    """Clone template_para, set its text, insert it directly after template_para. Returns new Paragraph."""
    from docx.text.paragraph import Paragraph
    clone = deepcopy(template_para._element)
    template_para._element.addnext(clone)
    p_new = Paragraph(clone, template_para._parent)
    set_text(p_new, new_text)
    return p_new


def insert_paragraphs_after(anchor_para, items):
    """
    items is a list of (style_name, text) pairs. Inserts in order after anchor_para.
    style_name should be a style that already exists in the doc (e.g. 'Heading 1', 'Heading 2',
    'Heading 3', 'normal'). For lists, use 'normal' and the caller can prefix bullets manually
    OR we apply numPr by cloning an existing list paragraph.
    """
    from docx.text.paragraph import Paragraph
    last_elem = anchor_para._element
    new_paras = []
    for style_name, text in items:
        # Clone anchor element so we inherit pPr (style/list-numbering); we'll override style below.
        clone = deepcopy(anchor_para._element)
        last_elem.addnext(clone)
        p_new = Paragraph(clone, anchor_para._parent)
        # Strip all runs
        for r in list(p_new.runs):
            r._element.getparent().remove(r._element)
        # Apply style by name (python-docx accepts the style name string directly)
        if style_name:
            try:
                p_new.style = style_name
            except KeyError:
                pass
        # Add the new text in a fresh run
        if text:
            p_new.add_run(text)
        new_paras.append(p_new)
        last_elem = clone
    return new_paras


def delete_para(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


# ---------- main ----------

def main():
    d = Document(SRC)
    paras = d.paragraphs
    tables = d.tables

    # 1. Cover page Phase descriptors (paras 14 and 15)
    set_text(paras[14], "PHASE 1   ·   DIY Foundation   ·   60 Days")
    set_text(paras[15], "PHASE 2   ·   Hudle Setup + Ads   ·   $1,000/month total (ads + consult)")

    # 2. "How to use this document" Phase 2 paragraph (para 28)
    set_text(
        paras[28],
        "Phase 2 (run concurrently or consecutively) brings Hudle Consulting in for a focused "
        "month-1 setup — Google Business Profile optimization, Google Local Services Ads and "
        "Search Ads campaign build, three Squarespace city landing pages, ad creative design, "
        "and a written handoff playbook plus curated training videos — followed by months 2–6 "
        "of bi-weekly check-in consults, ad creative refreshes, and progressively more city "
        "pages. Total $1,000/month covering BOTH ad spend (~$600–$700) AND the Hudle consult "
        "(~$300–$400). You manage the ads day-to-day with the playbook; Hudle is the strategist "
        "on call, not the ongoing ad manager."
    )

    # 3. Phase 1 cost/time/outcome paragraph (para 37) — drop the Squarespace upgrade cost
    set_text(
        paras[37],
        "Total cost: $0 (uses your existing Squarespace site as-is). Total time: ~12 hours "
        "focused work spread across 60 days. Expected outcome: First Google calls within 30–45 "
        "days; first AI citations within 60–90 days; meaningful Map Pack visibility in Saratoga "
        "Springs by day 60."
    )

    # 4. Budget table T2 — restructure to ads + consult model
    bt = tables[2]
    # Header row stays
    set_cell_text(bt.rows[1].cells[0], "Ad spend")
    set_cell_text(bt.rows[1].cells[1], "$600–$700")
    set_cell_text(bt.rows[1].cells[2], "Google Local Services Ads + Google Search Ads (uses $1,000 new-advertiser credit in month 1 if eligible)")
    set_cell_text(bt.rows[2].cells[0], "Hudle strategy consult")
    set_cell_text(bt.rows[2].cells[1], "$300–$400")
    set_cell_text(bt.rows[2].cells[2], "Month 1 setup + bi-weekly 30-min check-ins thereafter, ad creative refreshes, progressive landing pages")
    set_cell_text(bt.rows[3].cells[0], "Squarespace")
    set_cell_text(bt.rows[3].cells[1], "$0 (existing)")
    set_cell_text(bt.rows[3].cells[2], "Client's existing subscription — not a project cost")
    set_cell_text(bt.rows[4].cells[0], "Total")
    set_cell_text(bt.rows[4].cells[1], "$1,000")
    set_cell_text(bt.rows[4].cells[2], "Within budget; covers ads AND consult")

    # 5. GBP categories — primary already 'Landscape designer' (para 50); update secondary list (59-62)
    # New secondaries: Paving contractor, Landscaper, Landscape lighting designer, Masonry contractor,
    # Fireplace manufacturer, Lawn irrigation equipment supplier
    set_text(paras[59], "Paving contractor")
    set_text(paras[60], "Landscaper")
    set_text(paras[61], "Landscape lighting designer")
    set_text(paras[62], "Masonry contractor")
    # Insert two more list items after para 62, plus the conditional note
    inserted = insert_paragraphs_after(
        paras[62],
        [
            (paras[62].style.name, "Fireplace manufacturer (closest match for outdoor fireplaces)"),
            (paras[62].style.name, "Lawn irrigation equipment supplier"),
        ],
    )
    # Now append a conditional-note paragraph after the last inserted list item.
    # We want body text style, not list style — use 'normal'.
    insert_paragraphs_after(
        inserted[-1],
        [
            (
                "normal",
                "Conditional secondary — Landscape architect: only add this category if Aspen Stone "
                "holds a Utah DOPL Landscape Architect license (separate from the contractor "
                "license). If unsure, leave it off; misrepresenting a regulated license is a GBP "
                "suspension risk.",
            ),
        ],
    )

    # 6. Rewrite PHASE 2 section. Locate fresh paragraphs since indices may have shifted.
    paras = d.paragraphs  # refresh

    def find_idx(needle, start=0):
        for i in range(start, len(paras)):
            if needle in paras[i].text:
                return i
        raise ValueError(f"Not found: {needle!r}")

    # Cover/header in Phase 2 block
    i_phase2_title = find_idx("PHASE 2", start=200)
    i_phase2_sub = i_phase2_title + 1  # 'Scale-Up with Hudle Consulting ($1,000/month)'
    set_text(paras[i_phase2_sub], "Setup + Ads Engagement ($1,000/month total: ads + consult)")

    # Phase 2 intro paragraph
    i_phase2_intro = find_idx("Phase 2 starts when Phase 1", start=i_phase2_title)
    set_text(
        paras[i_phase2_intro],
        "Phase 2 starts when Phase 1's foundation is in place. The model is different from a "
        "traditional retainer: month 1 is concentrated setup work — Google Business Profile "
        "optimization, Google Local Services Ads and Google Search Ads campaign build, three "
        "Squarespace city landing pages, ad creative, and a complete handoff package — and "
        "months 2–6 shift to bi-weekly 30-minute check-ins, ad creative refreshes, monthly "
        "performance reviews, and additional landing pages built progressively. You run the "
        "ads day-to-day with the written playbook and curated training videos Hudle provides. "
        "Hudle is your on-call strategist, not your ad manager — that keeps the engagement at "
        "$1,000/month total (ads AND consult) instead of a typical $2k–$3k/month agency retainer."
    )

    # Section heading "What I do at $1,000/month" -> rename
    i_what_i_do = find_idx("What I do at $1,000/month")
    set_text(paras[i_what_i_do], "What Hudle does in Phase 2")

    # 'SEO work' subhead -> 'Month 1: One-time setup (intensive)'
    i_seo = find_idx("SEO work", start=i_what_i_do)
    set_text(paras[i_seo], "Month 1 — One-time setup (intensive)")

    # The four 'SEO work' bullets (i_seo+1 .. i_seo+4) — rewrite as setup bullets
    set_text(paras[i_seo + 1], "Google Business Profile optimization — categories audit (primary + secondaries per Phase 1 guidance), photo geotagging review, services and description finalization")
    set_text(paras[i_seo + 2], "Google Local Services Ads (LSAs) setup — DOPL license and insurance verification, profile build, lead-targeting cities and services configured")
    set_text(paras[i_seo + 3], "Google Search Ads campaign build — keyword targeting (paver patio + city, etc.), ad copy, location targeting, conversion tracking, $1,000 new-advertiser credit activation if eligible")
    set_text(paras[i_seo + 4], "Three Squarespace city landing pages (P1 cities: Saratoga Springs, Eagle Mountain, Lehi) — designed for Google Ads quality score AND organic ranking")

    # 'GEO work (the differentiator)' subhead -> 'Month 1: One-time setup (continued)'
    i_geo = find_idx("GEO work", start=i_seo)
    set_text(paras[i_geo], "Month 1 — One-time setup (continued)")

    # GEO work bullets (i_geo+1 .. i_geo+8 originally 8 bullets but let's confirm count)
    # Original bullets indices were i_geo+1..i_geo+8 (lines 249-256 in original). Rewrite all 8.
    setup_bullets_2 = [
        "Ad creative design — image and copy variants for LSAs and Search Ads, Utah County/hardscape-specific positioning",
        "Client training materials — written ads-monitoring playbook (daily, weekly, monthly checklists) plus a curated YouTube tutorial playlist (Surfside PPC, Define Digital Academy, Ben Heath)",
        "Handoff documentation — login inventory, escalation rules, what-to-watch-for guide, when-to-pause guide, monthly reporting template",
        "FAQ schema (JSON-LD) on the Phase 1 homepage FAQ — light-touch GEO add that boosts AI citation rates",
        "GEO baseline audit — run the prompt set across ChatGPT, Perplexity, and Google AI Overviews, document where Aspen Stone shows up vs. competitors",
        "First strategic Quora answer + Reddit ghost-write to seed authority",
        "Bing Webmaster Tools and Google Search Console health check",
        "30-day kickoff call + dedicated Slack/text channel for questions during setup",
    ]
    for offset, text in enumerate(setup_bullets_2, start=1):
        set_text(paras[i_geo + offset], text)

    # 'Reporting and strategy' subhead -> 'Months 2–6: Ongoing engagement'
    i_rs = find_idx("Reporting and strategy", start=i_geo)
    set_text(paras[i_rs], "Months 2–6 — Ongoing engagement")

    # Original 'Reporting and strategy' had 3 bullets. We want more — rewrite the 3 in place and append.
    ongoing_bullets = [
        "Bi-weekly 30-minute check-in calls — ad performance review, creative tweaks, troubleshooting, strategic questions",
        "Monthly performance report — leads, cost-per-lead, Map Pack rank, GBP calls, GEO citations, recommendations",
        "Ad creative refresh every 4–6 weeks — new image/copy variants to fight ad fatigue and improve quality score",
    ]
    set_text(paras[i_rs + 1], ongoing_bullets[0])
    set_text(paras[i_rs + 2], ongoing_bullets[1])
    set_text(paras[i_rs + 3], ongoing_bullets[2])

    # Insert additional ongoing bullets after paras[i_rs + 3]
    paras = d.paragraphs  # refresh
    i_rs = find_idx("Months 2–6 — Ongoing engagement")
    anchor_bullet = paras[i_rs + 3]
    extra_ongoing = [
        "Progressive city landing pages — Draper, Herriman, Highland (months 2–3); Cedar Hills, Spanish Fork, Mapleton (months 4–5)",
        "One GEO-optimized blog post per month — direct-answer format, FAQ schema, supports both organic and AI citations",
        "Continued Reddit/Quora ghost-writes — 2–3 deeper responses per month for you to post in your voice",
        "Monthly GEO visibility audit and citation tracking",
        "On-call advisor by text/email — you stay in the driver's seat on day-to-day ad management",
    ]
    insert_paragraphs_after(
        anchor_bullet,
        [(anchor_bullet.style.name, t) for t in extra_ongoing],
    )

    # 7. Phase 2 priorities by month — rewrite the four callout-box tables (T43..T46)
    # Refresh table refs (table count unchanged)
    tables = d.tables
    set_cell_text(
        tables[43].rows[0].cells[0],
        "Month 1 — Setup sprint\n"
        "Audit and lock GBP categories (primary: Landscape designer; six secondaries; conditional Landscape architect if licensed). "
        "Build and submit Google Local Services Ads profile (requires DOPL + insurance — Aspen Stone qualifies). "
        "Build Google Search Ads campaign and activate $1,000 new-advertiser credit if eligible. "
        "Design ad creative (image + copy variants). "
        "Build three P1 city landing pages on Squarespace: Saratoga Springs, Eagle Mountain, Lehi. "
        "Deliver written ads-monitoring playbook + curated YouTube training playlist. Kickoff call + ongoing text/email channel opens."
    )
    set_cell_text(
        tables[44].rows[0].cells[0],
        "Month 2\n"
        "Bi-weekly 30-min check-ins begin. First monthly performance report (leads, cost-per-lead, Map Pack rank, GEO citations). "
        "Begin P2 city landing pages: Draper and Herriman. First ad creative refresh based on what's converting. "
        "One GEO-optimized blog post: \"What a Paver Patio Costs in Utah County: 2026 Pricing Guide.\""
    )
    set_cell_text(
        tables[45].rows[0].cells[0],
        "Months 3–4\n"
        "Continue bi-weekly check-ins. Finish P2 city pages: Highland. Begin P3 city pages: Cedar Hills and Spanish Fork. "
        "Second ad creative refresh; expand search-keyword set if budget allows. "
        "GEO blog posts: \"Pavers vs. Concrete in Utah's Freeze-Thaw Climate\" and \"How to Choose a Hardscape Contractor in Utah County.\" "
        "Strategic Quora answers + Reddit ghost-writes."
    )
    set_cell_text(
        tables[46].rows[0].cells[0],
        "Months 5–6\n"
        "Finish P3 city pages: Mapleton. Third ad creative refresh. "
        "Mid-engagement performance review — what's working in ads, in landing pages, in GEO. "
        "Decide whether to expand ad budget, add Performance Max, or layer in retargeting based on results. "
        "Year-2 roadmap: which channels to keep, scale, or sunset."
    )

    # 8. Insert two new H1 sections after the "Phase 2 priorities by month" tables:
    # (a) Squarespace city landing page strategy, (b) Ad strategy.
    # The anchor is the last "month" callout table (T46). The next paragraph after that table
    # is currently the "The Nine-City Strategy" H1.
    # We need to insert paragraphs BEFORE "The Nine-City Strategy" H1.
    paras = d.paragraphs
    i_nine = find_idx("The Nine-City Strategy")

    # We'll insert above para[i_nine] by inserting after para[i_nine - 1].
    anchor = paras[i_nine - 1]

    # Build the two new sections as a list of (style, text). H1 + body + H3 sub-blocks.
    new_section_items = [
        ("Heading 1", "Squarespace City Landing Page Strategy"),
        ("normal", "City landing pages do double duty: they're the destinations for Google Ads (where ad quality score depends on landing-page relevance) AND they're the structure Google rewards for organic city-specific rankings. We build them on the existing Squarespace site — no platform migration."),
        ("Heading 3", "Build schedule"),
        ("normal", "Month 1 — P1 cities (ad destinations): Saratoga Springs, Eagle Mountain, Lehi."),
        ("normal", "Months 2–3 — P2 cities: Draper, Herriman, Highland."),
        ("normal", "Months 4–5 — P3 cities: Cedar Hills, Spanish Fork, Mapleton."),
        ("Heading 3", "Page structure (applied to every city)"),
        ("normal", "Headline with city + service (e.g., \"Paver Patio Installation in Lehi, UT\")."),
        ("normal", "Two to four project photos from that city when available; nearest-city photos when not."),
        ("normal", "One clear primary CTA — phone call or estimate form (not both competing for attention)."),
        ("normal", "Location-specific copy covering local context (soil, HOA quirks, climate, lot styles)."),
        ("normal", "Three to five FAQs in the GEO-friendly direct-answer format from Phase 1."),
        ("normal", "Contact form, NAP block, and DOPL license # at the bottom."),
        ("normal", "Each page is genuinely unique — Google penalizes near-duplicate templated city pages, and ad quality score does too."),
        ("Heading 1", "Ad Strategy"),
        ("normal", "We run two ad products in parallel because they serve different intent. Hudle designs the campaigns and creative; you operate them day-to-day with the written playbook and curated YouTube training."),
        ("Heading 3", "Google Local Services Ads (LSAs)"),
        ("normal", "Pay-per-lead (not per-click), shows above the Map Pack, and carries the Google Guaranteed badge. Requires DOPL license and insurance — Aspen Stone qualifies. This is the highest-trust placement on the result page for local service searches and should be the first product live."),
        ("Heading 3", "Google Search Ads"),
        ("normal", "Targets high-intent keywords like \"paver patio Lehi\" and \"hardscape contractor Saratoga Springs.\" If eligible, activate the $1,000 new-advertiser credit visible in the GBP dashboard during month 1 to extend reach without spending the budget on it. Search Ads route traffic to the city landing pages we just built."),
        ("Heading 3", "Ad creative"),
        ("normal", "Hudle designs image and copy variants tuned to the Utah County hardscape market — emphasizing the things that actually matter to local buyers (freeze-thaw durability, DOPL-licensed, real project photos, free estimates). Refreshed every 4–6 weeks to fight ad fatigue."),
        ("Heading 3", "Client handoff and operating model"),
        ("normal", "You manage the ads operationally after month-1 setup. You get: a written monitoring playbook (daily, weekly, monthly checks), a curated YouTube tutorial playlist (recommend: Surfside PPC, Define Digital Academy, Ben Heath), and bi-weekly 30-minute check-in calls with Hudle for questions, troubleshooting, and strategic decisions. Hudle is NOT the ongoing ad manager — that keeps the engagement at $1,000/month total."),
        ("Heading 3", "Bi-weekly cadence (months 2–6)"),
        ("normal", "Every other week: 30 minutes by phone or video. Topics: ad performance, what creative is winning, what budget shifts to make, lead quality, landing-page conversion, anything that came up. Between calls, async text/email for blockers."),
    ]
    insert_paragraphs_after(anchor, new_section_items)

    # 9. Update "What This Plan Intentionally Skips" section. Refresh.
    paras = d.paragraphs
    # Find the bullets to update
    i_skips = find_idx("What This Plan Intentionally Skips")
    # The current first sub-head is "No Google Ads or paid lead spend" — that's now FALSE.
    i_no_ads = find_idx("No Google Ads", start=i_skips)
    set_text(paras[i_no_ads], "Why we ARE running ads (this used to be skipped)")
    set_text(
        paras[i_no_ads + 1],
        "Earlier versions of this plan skipped paid ads. Two things changed: (1) Aspen Stone "
        "carries the DOPL license and insurance that qualify for Google Local Services Ads — "
        "the most efficient lead product Google offers, paid per lead, not per click; and "
        "(2) the $1,000 new-advertiser Google Ads credit visible in the GBP dashboard extends "
        "month-1 reach essentially for free. We've shifted ~$600–$700/month of the budget to "
        "ads and kept the rest for the Hudle consult instead of a full agency retainer."
    )
    # "No agency retainers" — keep but tweak slightly
    i_no_retainer = find_idx("No agency retainers", start=i_skips)
    set_text(
        paras[i_no_retainer + 1],
        "Most local SEO agencies charge $1,500–$3,000/month and don't do GEO at all yet. "
        "This setup-plus-consult model gives you ad infrastructure + landing pages + GEO + "
        "bi-weekly strategy access at $1,000/month total — because Hudle hands you the keys "
        "to operate the ads, instead of running them on a retainer."
    )

    # 10. Risks section — update the "Competitor outspends with Google Ads" risk
    paras = d.paragraphs
    i_risks = find_idx("Risks & Mitigations")
    i_comp = find_idx("Risk: Competitor outspends", start=i_risks)
    set_text(paras[i_comp], "Risk: Ad spend gets wasted on low-quality leads")
    set_text(
        paras[i_comp + 1],
        "LSAs and Search Ads both produce some unqualified contacts (price shoppers, out-of-area, "
        "DIYers). Mitigation: tight geo-targeting to the nine cities, lead-quality review every "
        "bi-weekly check-in, dispute unqualified LSA leads inside Google's window (they refund). "
        "Pause underperforming keywords monthly. The playbook tells you exactly when to escalate "
        "to Hudle vs. handle yourself."
    )

    # Squarespace platform limitations risk — soften since we de-emphasize schema
    i_sq = find_idx("Risk: Squarespace platform", start=i_risks)
    set_text(
        paras[i_sq + 1],
        "Squarespace is workable for the city landing pages, FAQ section, and ads destinations "
        "we need — that's the bulk of the work. Advanced schema (beyond LocalBusiness and FAQ) "
        "and root-level llms.txt are harder; we treat those as nice-to-have, not core. Revisit a "
        "WordPress migration only if organic plateaus after month 9–12."
    )

    # Save
    d.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
